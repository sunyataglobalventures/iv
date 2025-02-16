from flask import Flask, request, render_template, send_file, redirect, url_for
from google.cloud import firestore
import os
import json
import base64
from docx import Document
from datetime import datetime
import firebase_admin
from firebase_admin import credentials, firestore

app = Flask(__name__)


# # Set the Firebase service account key
# os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "credentials.json"

# # Initialize Firestore client
# db = firestore.Client()

# Initialize Firebase
firebase_key_base64 = os.getenv("FIREBASE_KEY")
if not firebase_key_base64:
    raise ValueError("FIREBASE_KEY environment variable is not set")

firebase_key = json.loads(base64.b64decode(firebase_key_base64).decode("utf-8"))
cred = credentials.Certificate(firebase_key)
firebase_admin.initialize_app(cred)

db = firestore.client()

# Helper functions
from docx.shared import Pt  # Import for text formatting

def replace_text_in_run(run, key, value):
    """Replaces text in a run and makes it bold."""
    if key in run.text:
        run.text = run.text.replace(key, value)
        run.bold = True  # Make text bold
        run.font.size = Pt(12)  # Optional: Set font size

def replace_placeholders(doc, placeholders):
    """Replaces placeholders in paragraphs and tables."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in placeholders.items():
                replace_text_in_run(run, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in placeholders.items():
                            replace_text_in_run(run, key, value)



def get_template_path(invoice_type):
    """Select the template based on invoice type."""
    return "INVOICE.docx" if invoice_type == "invoice" else "PROFORMA_INVOICE.docx"

def create_invoice(data, template_path, output_folder, unique_id):
    """Generate invoice document by replacing placeholders."""
    doc = Document(template_path)

    placeholders = {
        "[IVN]": data.get("invoice_no", "N/A"),
        "[DAT]": data.get("invoice_date", "N/A"),
        "[IDD]": data.get("due_date", "N/A"),
        "[NAME]": data.get("name", "N/A"),
        "[STORENAME]": data.get("store_name", "N/A"),
        "[ADDRESS]": data.get("address", "N/A"),
        "[PHN]": data.get("phone", "N/A"),
        "[EMAIL]": data.get("email", "N/A"),
        "[SERVICE]": data.get("service", "N/A"),
        "[COST]": data.get("cost", "N/A"),
        "[GT]": data.get("gst", "N/A"),
        "MRP": data.get("total", "N/A"),
    }

    replace_placeholders(doc, placeholders)

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    file_name = f"{data.get('invoice_type', 'INVOICE')}_{data.get('service', 'N/A')}_{data.get('store_name', 'N/A')}_{data.get('invoice_date', 'N/A')}.docx"
    file_path = os.path.join(output_folder, file_name)
    doc.save(file_path)
    
    return file_path

def save_invoice_to_firestore(data):
    """Save invoice data in Firestore under 'INVOICES' collection."""
    collection_name = "INVOICES"
    data["timestamp"] = datetime.utcnow().isoformat()

    doc_ref = db.collection(collection_name).document()
    unique_id = doc_ref.id
    data["unique_id"] = unique_id
    doc_ref.set(data)
    
    return unique_id

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        data = request.form.to_dict()

        invoice_type = data.get("invoice_type")
        template_path = get_template_path(invoice_type)

        output_folder = "invoices"

        try:
            unique_id = save_invoice_to_firestore(data)
            file_path = create_invoice(data, template_path, output_folder, unique_id)

            return redirect(url_for("download", file_name=os.path.basename(file_path)))
        except Exception as e:
            return f"An error occurred: {e}", 500

    return render_template("index.html")

@app.route("/download/<file_name>")
def download(file_name):
    file_path = os.path.join("invoices", file_name)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return "File not found", 404

if __name__ == "__main__":
    app.run(debug=True)
