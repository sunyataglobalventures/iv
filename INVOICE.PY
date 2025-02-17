import os
from openpyxl import load_workbook
from docx import Document
from datetime import datetime

def replace_text_in_run(run, key, value):
    if key in run.text:
        run.text = run.text.replace(key, value)
        run.font.bold = True

def replace_placeholders(doc, placeholders):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in placeholders.items():
                replace_text_in_run(run, key, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in placeholders.items():
                            replace_text_in_run(run, key, value)

def create_offer_letter(template_path, excel_path, output_folder, serial_number=None):
    # Load Excel workbook
    wb = load_workbook(excel_path)
    ws = wb.active
    
    # Load offer letter template
    doc = Document(template_path)
    
    if serial_number:
        row = ws[serial_number]
        invoice, ivdate ,duedate, name, storename, address, phone, email,service,cost,gst, total = [cell.value for cell in row]
        
         # Ensure the date is formatted correctly for the file name
        if isinstance(ivdate, datetime):
            formatted_date = ivdate.strftime('%Y-%m-%d')
        else:
            # If date is a string, ensure it's formatted properly
            formatted_date = ivdate.split()[0].replace('/', '-').replace(':', '-')
        
        # Construct file name
        file_name = f"INVOICE_{service}_{storename}_{formatted_date}.docx"
        file_path = os.path.join(output_folder, file_name)
        
       
        print("Current working directory:", os.getcwd())
        print("File path:", file_path)
        
        # Create output folder if it doesn't exist
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
        
        # Check if file already exists
        if os.path.exists(file_path):
            print(f"File '{file_name}' already exists.")
        else:

            

            # Prepare placeholders
            placeholders = {
               
                "[IVN]": str(invoice),  
                "[DAT]": ivdate.strftime("%d/%m/%Y") if isinstance(ivdate, datetime) else str(ivdate),
                "[IDD]": duedate.strftime("%d/%m/%Y") if isinstance(duedate, datetime) else str(duedate),
                "[NAME]": str(name),                               
                "[STORENAME]": str(storename), 
                "[ADDRESS]": str(address),
                "[PHN]": str(phone), 
                "[EMAIL]": str(email),            
                "[SERVICE]": str(service),
                "[COST]": str(cost) ,
                "[GT]": str(gst),
                "MRP": str(total),  
                
            }

            print("Placeholders being replaced:")
            for key, value in placeholders.items():
                print(f"{key}: {value}")
            
            # Replace placeholders in the document
            replace_placeholders(doc, placeholders)
            
            # Save new offer letter
            doc.save(file_path)
            print(f"INVOICE letter created: '{file_name}'")
            
            # Check if file was successfully created
            if os.path.exists(file_path):
                print(f"File '{file_name}' created successfully.")
            else:
                print(f"Failed to create file '{file_name}'.")
    else:
        print("Please enter a serial number.")

if __name__ == "__main__":
    template_path = "INVOICE.docx"  # Path to your offer letter template
    excel_path = "INVOICE.xlsx"  # Path to your Excel sheet
    output_folder = "INVOICE"  # Folder where offer letters will be saved

    while True:
        serial_number = input("Enter the serial number for the employee (or type 'exit' to quit): ")
        if serial_number.lower() == 'exit':
            break
        try:
            create_offer_letter(template_path, excel_path, output_folder, serial_number)
        except Exception as e:
            print(f"An error occurred: {e}")
